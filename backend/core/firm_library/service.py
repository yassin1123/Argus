"""Firm-library ingestion + retirement service.

Reuses the existing chunking dispatch in ``backend/ingest/chunking`` (PDF
chunker for PDFs, web/knowledge chunker for plain text). DOCX is the
one new path: parsed via ``python-docx`` to plain text, then handed to
the knowledge chunker so the rest of the pipeline (embeddings, insert,
retrieval, NLI) is shared.

Idempotency is keyed on ``(firm_id, sha256(file_bytes))``: re-running
the same upload returns the existing row without re-embedding. Useful
when a partner re-runs a build script or sync job.

Retire is a soft-delete: the ``firm_content`` row gets ``retired_at``
set, and every chunk linked to it gets a ``metadata->>'retired_at'``
stamp so the retrieval-side filter excludes them. We deliberately don't
delete the chunk rows — historical engagement citations to a retired
playbook stay valid (the writer can still surface a quote with a
"retired" badge), only NEW retrieval excludes them.
"""

from __future__ import annotations

import hashlib
import io
import logging
from dataclasses import dataclass, field
from typing import Any, Literal

from core.embeddings import embed_texts
from ingest.chunking import Chunk, chunk_source
from storage.chunk_queries import insert_chunks
from storage.firm_content_queries import (
    find_active_by_filehash,
    insert_firm_content,
    list_chunks_for_content,
    mark_chunks_retired,
    retire_firm_content_row,
    update_chunk_count,
)

logger = logging.getLogger(__name__)

FIRM_LIBRARY_SOURCE_TYPE = "firm_library"

Category = Literal[
    "playbook", "sector_primer", "prior_report", "framework", "methodology", "other"
]

# File extensions we handle today. New types should land in the existing
# chunking dispatch (ingest/chunking/__init__.py) so this list extends
# naturally.
SUPPORTED_EXTENSIONS: tuple[str, ...] = (".pdf", ".docx", ".md", ".txt")


class UnsupportedFileTypeError(ValueError):
    """Raised when the uploaded file extension isn't in SUPPORTED_EXTENSIONS."""


@dataclass
class IngestResult:
    """Return shape of :func:`ingest_firm_content`."""

    firm_content_id: str
    cached: bool = False
    chunks_written: int = 0
    chunk_ids: list[str] = field(default_factory=list)


def _file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _docx_to_text(file_bytes: bytes) -> str:
    """Extract text from a .docx using python-docx (in deps as 1.1.2)."""
    import docx  # type: ignore  # noqa: WPS433

    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # Tables aren't structurally rich here — flatten to "row | row | row"
    # so the chunker still gets the content.
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _chunk_uploaded_bytes(
    *, file_bytes: bytes, source_filename: str
) -> tuple[list[Chunk], str]:
    """Pick the right chunker for the file type. Returns ``(chunks, chunker_kind)``.

    The chunker_kind is logged so we have provenance on which path
    produced the chunks.
    """
    name = (source_filename or "").lower()
    if name.endswith(".pdf"):
        chunks = chunk_source(source_kind="pdf", content=file_bytes)
        return chunks, "pdf"
    if name.endswith(".docx"):
        text = _docx_to_text(file_bytes)
        chunks = chunk_source(source_kind="knowledge", content=text)
        return chunks, "docx"
    if name.endswith(".md") or name.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="replace")
        chunks = chunk_source(source_kind="knowledge", content=text)
        return chunks, "text"
    raise UnsupportedFileTypeError(
        f"file extension not supported: {source_filename!r}. "
        f"Supported: {SUPPORTED_EXTENSIONS}"
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


async def ingest_firm_content(
    *,
    firm_id: str,
    title: str,
    category: Category,
    file_bytes: bytes,
    source_filename: str,
    uploaded_by: str | None,
    description: str | None = None,
    intended_modes: list[str] | None = None,
    sector_tags: list[str] | None = None,
    trust_level: str = "firm_vetted",
) -> IngestResult:
    """Parse → chunk → embed → write firm-scoped chunks.

    Idempotent on ``(firm_id, sha256(file_bytes))``: a second call with
    identical bytes returns the existing record without re-embedding.

    Raises :class:`UnsupportedFileTypeError` for non-allowlist extensions.
    Raises ``ValueError`` for empty bodies / chunkers producing zero chunks.
    """
    if not file_bytes:
        raise ValueError("empty file body")

    fh = _file_hash(file_bytes)
    cached = await find_active_by_filehash(firm_id, fh)
    if cached:
        logger.info(
            "firm_library: cache HIT firm=%s file_hash=%s — returning %s",
            firm_id,
            fh[:12],
            cached["id"],
        )
        return IngestResult(
            firm_content_id=str(cached["id"]),
            cached=True,
            chunks_written=int(cached.get("chunk_count") or 0),
        )

    # Parse + chunk first; if the file is unparseable we don't want a
    # half-written firm_content row hanging around.
    chunks, chunker_kind = _chunk_uploaded_bytes(
        file_bytes=file_bytes, source_filename=source_filename
    )
    if not chunks:
        raise ValueError(
            f"chunker produced 0 chunks for {source_filename!r} "
            f"(parser={chunker_kind}). File may be empty or scanned-only."
        )

    contents = [c.content for c in chunks]
    embeddings = await embed_texts(contents)
    if len(embeddings) != len(chunks):
        raise RuntimeError(
            f"embedding count mismatch: got {len(embeddings)} for {len(chunks)} chunks"
        )

    # Write the firm_content row first so we have an id to thread through
    # into chunk metadata.
    fc = await insert_firm_content(
        firm_id=firm_id,
        title=title,
        category=category,
        description=description,
        intended_modes=intended_modes or [],
        sector_tags=sector_tags or [],
        source_filename=source_filename,
        file_hash=fh,
        trust_level=trust_level,
        uploaded_by=uploaded_by,
        metadata={"chunker": chunker_kind},
    )
    fc_id = str(fc["id"])

    rows: list[dict[str, Any]] = []
    for c, vec in zip(chunks, embeddings):
        rows.append(
            {
                "content": c.content,
                "content_hash": c.content_hash,
                "embedding": vec,
                "position": c.position,
                "page": c.page,
                "slide": c.slide,
                "timestamp_str": c.timestamp_str,
                "speaker": c.speaker,
                "section_heading": c.section_heading,
                "metadata": {
                    "firm_content_id": fc_id,
                    "category": category,
                    "intended_modes": list(intended_modes or []),
                    "sector_tags": list(sector_tags or []),
                    "title": title,
                },
            }
        )

    chunk_ids = await insert_chunks(
        session_id=None,                # firm-global — visible to every engagement at this firm
        blob_id=None,
        source_file_id=None,
        source_type=FIRM_LIBRARY_SOURCE_TYPE,
        source_filename=source_filename,
        source_url=None,
        trust_level=trust_level,
        rows=rows,
        firm_id=firm_id,
        firm_content_id=fc_id,
    )
    await update_chunk_count(fc_id, len(chunk_ids), absolute=True)

    return IngestResult(
        firm_content_id=fc_id,
        cached=False,
        chunks_written=len(chunk_ids),
        chunk_ids=chunk_ids,
    )


async def retire_firm_content(
    *,
    firm_id: str,
    content_id: str,
    retired_by: str | None,
) -> dict[str, Any] | None:
    """Soft-delete a firm_content row + stamp its chunks for retrieval-side
    filtering. Returns the updated firm_content dict, or ``None`` when the
    row was already retired or doesn't belong to ``firm_id``.

    Audit-log entry is written so the retire action shows up in
    compliance reports. Caller is responsible for HTTP-level audit too
    if they want method/path captured.
    """
    fc = await retire_firm_content_row(firm_id, content_id, retired_by)
    if not fc:
        return None
    n_chunks = await mark_chunks_retired(firm_id, content_id)
    logger.info(
        "firm_library: retired %s (firm=%s, chunks_marked=%d)",
        content_id,
        firm_id,
        n_chunks,
    )

    # Audit log — uses the existing audit_events table from migration 021.
    from db.connection import acquire  # noqa: WPS433

    async with acquire() as conn:
        await conn.execute(
            """
            INSERT INTO audit_events (
                actor_user_id, action, resource_type, resource_id, payload
            ) VALUES (
                $1::uuid, 'firm_library.retire', 'firm_content', $2, $3::jsonb
            )
            """,
            retired_by,
            content_id,
            f'{{"firm_id":"{firm_id}","chunks_marked":{n_chunks}}}',
        )
    return fc


# Re-export for convenience (callers reaching for the chunk preview in the
# GET-one endpoint).
__all__ = [
    "FIRM_LIBRARY_SOURCE_TYPE",
    "SUPPORTED_EXTENSIONS",
    "Category",
    "IngestResult",
    "UnsupportedFileTypeError",
    "ingest_firm_content",
    "list_chunks_for_content",
    "retire_firm_content",
]
