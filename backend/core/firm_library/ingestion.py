"""Firm-library batch ingestion + hardening — W14/D2.

Wraps the W5 single-file ``service.ingest_firm_content`` with batch
behaviour real firm uploads need:

  - **Error isolation per file** — one corrupt PDF doesn't fail the
    batch. Each file gets its own try/except and reports its outcome
    in :class:`IngestionResult`.
  - **Structured per-file results** — filename, status
    (``ready`` / ``failed`` / ``dedup_skipped``), chunks_created,
    error_reason, file_hash, firm_content_id. Callers (CLI, seeders,
    future API batch endpoint) walk the list and report cleanly.
  - **Content-type detector** — extension-driven routing to the right
    extractor. Today: PDF / DOCX / MD / TXT / CSV. xlsx surfaces a
    clear "not supported" reason rather than crashing.
  - **Sentence-aware chunking for text/markdown** — routes through
    :func:`chunker.chunk_library_text` which respects sentence
    boundaries, keeps tables whole, and prepends a small overlap so
    boundary-straddling claims stay retrievable.

The single-file API endpoint at ``backend/api/firm_library.py`` keeps
using ``service.ingest_firm_content`` unchanged — the batch path is
additive. New CLI / seeders should use :func:`ingest_directory` or
:func:`ingest_files`.
"""

from __future__ import annotations

import csv as _csv
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Literal

from core.embeddings import embed_texts
from ingest.chunking import Chunk
from storage.chunk_queries import insert_chunks
from storage.firm_content_queries import (
    find_active_by_filehash,
    insert_firm_content,
    update_chunk_count,
)

from .chunker import chunk_library_text
from .service import FIRM_LIBRARY_SOURCE_TYPE, _file_hash

logger = logging.getLogger(__name__)


Status = Literal["ready", "failed", "dedup_skipped"]

# Content-type routing — extension → (extractor_kind, reason if unsupported).
# Add new types here as extractors land. ``xlsx`` is documented as
# unsupported with a clear reason so callers know to convert to CSV.
_EXT_ROUTING: dict[str, str] = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".md":   "text",
    ".txt":  "text",
    ".csv":  "csv",
}

_UNSUPPORTED_EXTS: dict[str, str] = {
    ".xlsx": "Excel workbooks aren't supported yet — convert to CSV first (W14/D2 carry-forward).",
    ".xls":  "Excel workbooks aren't supported yet — convert to CSV first.",
    ".pptx": "PowerPoint decks aren't supported as a library source — convert to PDF first.",
    ".html": "HTML pages aren't supported via the library upload path — convert to PDF or paste as markdown.",
    ".rtf":  "RTF isn't supported — convert to docx or markdown first.",
}


@dataclass
class IngestionResult:
    """Per-file outcome for a batch run. Callers walk a list of these
    to surface user-facing summaries.

    ``status``:
      - ``ready``: parsed, chunked, embedded, persisted. ``chunks_created`` > 0.
      - ``dedup_skipped``: identical file_hash already exists for this firm.
        ``firm_content_id`` points at the existing row.
      - ``failed``: anything went wrong. ``error_reason`` is set to a
        short human-readable string.
    """

    filename: str
    status: Status
    chunks_created: int = 0
    error_reason: str | None = None
    dedup_skipped: bool = False
    file_hash: str | None = None
    firm_content_id: str | None = None
    extractor: str | None = None
    extra: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Content-type detection
# ---------------------------------------------------------------------------


def detect_content_type(filename: str) -> tuple[str | None, str | None]:
    """Map a filename to ``(extractor_kind, unsupported_reason)``.

    Exactly one of the two will be non-None. ``extractor_kind`` is the
    routing token used by :func:`_extract_text` (``pdf`` / ``docx`` /
    ``text`` / ``csv``). ``unsupported_reason`` is a human-readable
    string explaining why the extension isn't routed.
    """
    ext = Path(filename).suffix.lower()
    if ext in _EXT_ROUTING:
        return _EXT_ROUTING[ext], None
    if ext in _UNSUPPORTED_EXTS:
        return None, _UNSUPPORTED_EXTS[ext]
    if not ext:
        return None, "missing file extension — can't route to an extractor."
    return None, (
        f"unsupported file extension {ext!r}. "
        f"Supported: {sorted(_EXT_ROUTING.keys())}."
    )


# ---------------------------------------------------------------------------
# Extractors (one per content type)
# ---------------------------------------------------------------------------


def _extract_pdf(file_bytes: bytes) -> list[Chunk]:
    """PDF route — delegates to the existing W5 PDF chunker so
    page-numbered chunks land identically to the API path. We don't
    re-pack PDF chunks through the sentence-aware text chunker because
    page boundaries are themselves a natural and useful retrieval unit.
    """
    from ingest.chunking import chunk_source

    return chunk_source(source_kind="pdf", content=file_bytes)


def _extract_docx(file_bytes: bytes) -> list[Chunk]:
    """DOCX route — extract text + flatten tables via python-docx, then
    pack through the sentence-aware library chunker."""
    import docx  # type: ignore  # noqa: WPS433

    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append("\n".join(rows))
    text = "\n\n".join(parts).strip()
    return chunk_library_text(text)


def _extract_text(file_bytes: bytes) -> list[Chunk]:
    """Markdown / TXT route — UTF-8 decode (replace on bad bytes) and
    pack through the sentence-aware library chunker."""
    text = file_bytes.decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError("file decoded to an empty string")
    return chunk_library_text(text)


def _extract_csv(file_bytes: bytes) -> list[Chunk]:
    """CSV route — decode + walk rows. Each row becomes a single line
    formatted as a markdown table row so the library chunker can keep
    it whole. The header row is repeated at the top of every chunk
    (caller's responsibility to set a short CSV; we don't paginate
    headers across chunks for simplicity).
    """
    text = file_bytes.decode("utf-8", errors="replace")
    reader = _csv.reader(io.StringIO(text))
    rows = [r for r in reader if any((c or "").strip() for c in r)]
    if not rows:
        raise ValueError("CSV decoded to zero non-empty rows")
    header = rows[0]
    data = rows[1:] or []
    md_lines = ["| " + " | ".join(header) + " |"]
    md_lines.append("|" + "|".join(["---"] * len(header)) + "|")
    for row in data:
        cells = [(c or "").strip() for c in row]
        if len(cells) < len(header):
            cells = cells + [""] * (len(header) - len(cells))
        elif len(cells) > len(header):
            cells = cells[: len(header)]
        md_lines.append("| " + " | ".join(cells) + " |")
    text = "\n".join(md_lines)
    return chunk_library_text(text)


_EXTRACTORS = {
    "pdf":  _extract_pdf,
    "docx": _extract_docx,
    "text": _extract_text,
    "csv":  _extract_csv,
}


# ---------------------------------------------------------------------------
# Single-file ingestion (hardened) — used by the batch path
# ---------------------------------------------------------------------------


async def _ingest_single_hardened(
    *,
    firm_id: str,
    title: str,
    category: str,
    file_bytes: bytes,
    source_filename: str,
    uploaded_by: str | None,
    description: str | None,
    intended_modes: list[str],
    sector_tags: list[str],
    trust_level: str,
) -> IngestionResult:
    """Single-file path with structured-result reporting. Mirrors the
    shape of ``service.ingest_firm_content`` but never raises — any
    exception is captured into ``IngestionResult.error_reason``.
    """
    result = IngestionResult(
        filename=source_filename,
        status="failed",
    )

    if not file_bytes:
        result.error_reason = "empty file body"
        return result

    extractor_kind, unsupported = detect_content_type(source_filename)
    if unsupported:
        result.error_reason = unsupported
        return result
    if extractor_kind is None:
        result.error_reason = "no extractor registered"
        return result
    result.extractor = extractor_kind

    fh = _file_hash(file_bytes)
    result.file_hash = fh

    # Dedup against the firm's existing content.
    try:
        cached = await find_active_by_filehash(firm_id, fh)
    except Exception as e:  # noqa: BLE001
        result.error_reason = f"dedup lookup failed: {e}"
        return result
    if cached:
        result.status = "dedup_skipped"
        result.dedup_skipped = True
        result.firm_content_id = str(cached["id"])
        result.chunks_created = int(cached.get("chunk_count") or 0)
        return result

    # Parse → chunk.
    try:
        chunks = _EXTRACTORS[extractor_kind](file_bytes)
    except Exception as e:  # noqa: BLE001
        result.error_reason = f"{extractor_kind} extractor failed: {e}"
        return result
    if not chunks:
        result.error_reason = (
            f"{extractor_kind} extractor produced 0 chunks "
            "(file may be empty or scanned-only)"
        )
        return result

    # Embed.
    try:
        contents = [c.content for c in chunks]
        embeddings = await embed_texts(contents)
    except Exception as e:  # noqa: BLE001
        result.error_reason = f"embedding failed: {e}"
        return result
    if len(embeddings) != len(chunks):
        result.error_reason = (
            f"embedding count mismatch: got {len(embeddings)} for {len(chunks)} chunks"
        )
        return result

    # Persist firm_content row.
    try:
        fc = await insert_firm_content(
            firm_id=firm_id,
            title=title,
            category=category,
            description=description,
            intended_modes=intended_modes,
            sector_tags=sector_tags,
            source_filename=source_filename,
            file_hash=fh,
            trust_level=trust_level,
            uploaded_by=uploaded_by,
            metadata={"chunker": extractor_kind, "hardened": True},
        )
    except Exception as e:  # noqa: BLE001
        result.error_reason = f"firm_content insert failed: {e}"
        return result
    fc_id = str(fc["id"])
    result.firm_content_id = fc_id

    rows: list[dict[str, Any]] = []
    for c, vec in zip(chunks, embeddings):
        rows.append(
            {
                "content": c.content,
                "content_hash": c.content_hash,
                "embedding": vec,
                "position": c.position,
                "page": c.page,
                "slide": c.slide,
                "timestamp_str": c.timestamp_str,
                "speaker": c.speaker,
                "section_heading": c.section_heading,
                "metadata": {
                    "firm_content_id": fc_id,
                    "category": category,
                    "intended_modes": list(intended_modes),
                    "sector_tags": list(sector_tags),
                    "title": title,
                },
            }
        )

    try:
        chunk_ids = await insert_chunks(
            session_id=None,
            blob_id=None,
            source_file_id=None,
            source_type=FIRM_LIBRARY_SOURCE_TYPE,
            source_filename=source_filename,
            source_url=None,
            trust_level=trust_level,
            rows=rows,
            firm_id=firm_id,
            firm_content_id=fc_id,
        )
        await update_chunk_count(fc_id, len(chunk_ids), absolute=True)
    except Exception as e:  # noqa: BLE001
        result.error_reason = f"chunk persist failed: {e}"
        return result

    result.status = "ready"
    result.chunks_created = len(chunk_ids)
    return result


# ---------------------------------------------------------------------------
# Batch entry points
# ---------------------------------------------------------------------------


async def ingest_files(
    *,
    firm_id: str,
    files: Iterable[tuple[str, bytes]],
    category: str,
    intended_modes: list[str] | None = None,
    sector_tags: list[str] | None = None,
    trust_level: str = "firm_vetted",
    uploaded_by: str | None = None,
    title_fn: Any = None,
) -> list[IngestionResult]:
    """Bulk ingestion from an iterable of ``(filename, bytes)`` pairs.

    Each file is ingested independently; one failure doesn't abort
    the batch. Returns a list of :class:`IngestionResult` in the same
    order as the input.

    ``title_fn`` (optional) is a callable ``(filename, bytes) -> str``
    that picks the document title shown in the library UI. Defaults
    to the filename stem.
    """
    intended_modes = list(intended_modes or [])
    sector_tags = list(sector_tags or [])
    title_resolver = title_fn or (lambda fname, _b: Path(fname).stem.replace("_", " ").title())

    results: list[IngestionResult] = []
    for fname, body in files:
        try:
            title = str(title_resolver(fname, body))
        except Exception:
            title = Path(fname).stem
        res = await _ingest_single_hardened(
            firm_id=firm_id,
            title=title,
            category=category,
            file_bytes=body,
            source_filename=fname,
            uploaded_by=uploaded_by,
            description=None,
            intended_modes=intended_modes,
            sector_tags=sector_tags,
            trust_level=trust_level,
        )
        results.append(res)
    return results


async def ingest_directory(
    *,
    firm_id: str,
    directory: Path,
    category: str,
    intended_modes: list[str] | None = None,
    sector_tags: list[str] | None = None,
    trust_level: str = "firm_vetted",
    uploaded_by: str | None = None,
    recursive: bool = False,
) -> list[IngestionResult]:
    """Ingest every supported file in ``directory``. Order is the
    directory listing order (sorted by filename for determinism).
    Hidden files (leading ``.``) are skipped. Files with unsupported
    extensions land as ``failed`` rows with a clear ``error_reason``
    (see :func:`detect_content_type`).
    """
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"not a directory: {directory}")

    if recursive:
        paths = sorted(p for p in directory.rglob("*") if p.is_file() and not p.name.startswith("."))
    else:
        paths = sorted(p for p in directory.iterdir() if p.is_file() and not p.name.startswith("."))

    files = [(p.name, p.read_bytes()) for p in paths]
    return await ingest_files(
        firm_id=firm_id,
        files=files,
        category=category,
        intended_modes=intended_modes,
        sector_tags=sector_tags,
        trust_level=trust_level,
        uploaded_by=uploaded_by,
    )


def summarise(results: list[IngestionResult]) -> dict[str, Any]:
    """Roll up a batch's per-file outcomes into a single summary dict
    used by the CLI + seeder reports."""
    by_status: dict[str, int] = {"ready": 0, "dedup_skipped": 0, "failed": 0}
    total_chunks = 0
    failures: list[dict[str, str]] = []
    for r in results:
        by_status[r.status] = by_status.get(r.status, 0) + 1
        if r.status == "ready":
            total_chunks += r.chunks_created
        if r.status == "failed":
            failures.append({"filename": r.filename, "reason": r.error_reason or "unknown"})
    return {
        "total_files": len(results),
        "by_status": by_status,
        "chunks_created": total_chunks,
        "failures": failures,
    }


__all__ = [
    "IngestionResult",
    "Status",
    "detect_content_type",
    "ingest_directory",
    "ingest_files",
    "summarise",
]
